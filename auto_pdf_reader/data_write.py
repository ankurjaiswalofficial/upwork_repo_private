import docx
from docx.shared import Pt
import docx2txt

doc = docx.Document()

data = docx2txt.process("./auto_pdf_reader/Story.docx")
# def page_runner(i):
#     pr1 = doc.add_paragraph()
#     pr1.paragraph_format.line_spacing = 1.5
#     para = pr1.add_run()
#     para.font.size = Pt(14)
#     para.font.name = "Garamond"
#     with open(f"./auto_pdf_reader/data_files/data_file_{i}.txt", "r") as f:
#         para.text = f.read()
#     doc.add_page_break()

def page_runner(data):
    pr1 = doc.add_paragraph()
    pr1.paragraph_format.line_spacing = 1.5
    para = pr1.add_run()
    para.font.size = Pt(14)
    para.font.name = "Garamond"
    para.text = data


for i in range(0, 75):
    # f = open("work.txt", "w")
    # f.write(data)
    # f.close()
    page_runner(data)

doc.save('gfgv.docx')
